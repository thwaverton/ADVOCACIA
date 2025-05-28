# para iniciar, no terminal: streamlit run seu_arquivo_python.py
# (substitua seu_arquivo_python.py pelo nome real do arquivo)

import streamlit as st
import requests
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup, NavigableString
import sys 
import subprocess
import json
import os  # Importado para lidar com nomes de arquivo na transcrição
from pypdf import PdfReader  # <<< NOVO: Para ler PDFs
import toml  # Para verificar se o arquivo secrets.toml existe e tem as chaves (opcional, mas bom para feedback)

# --- Configurações Globais e Constantes ---
CHATVOLT_API_BASE_URL = "https://api.chatvolt.ai/agents"
GROQ_API_BASE_URL = "https://api.groq.com/openai/v1"
COMMON_GROQ_MODELS = ["llama3-8b-8192", "llama3-70b-8192", "mixtral-8x7b-32768", "gemma-7b-it"]

# Constantes para Transcrição com Groq
GROQ_API_TRANSCRIPTIONS_ENDPOINT = "https://api.groq.com/openai/v1/audio/transcriptions"
SELECTED_TRANSCRIPTION_MODEL = "whisper-large-v3-turbo"  # Mais rápido para transcrição PT
MAX_AUDIO_FILE_SIZE_MB = 25  # Limite da API Groq

# <<< NOVO: Constantes para upload de arquivos de texto >>>
ALLOWED_TEXT_EXTENSIONS = ["txt", "pdf", "docx"]


# --- Funções Utilitárias ---

# (Funções de DOCX: add_runs_from_html_element, create_docx_from_text_or_html permanecem as mesmas)
def add_runs_from_html_element(paragraph, element):
    for child in element.children:
        if isinstance(child, NavigableString):
            text_content = str(child)
            if text_content.strip():
                paragraph.add_run(text_content)
            elif text_content:  # Adiciona espaços em branco se eles existirem
                paragraph.add_run(text_content)
        elif child.name in ['strong', 'b']:
            run = paragraph.add_run()
            add_runs_from_html_element(run, child)
            run.bold = True
        elif child.name in ['em', 'i']:
            run = paragraph.add_run()
            add_runs_from_html_element(run, child)
            run.italic = True
        elif child.name == 'br':
            paragraph.add_run().add_break()
        elif child.name in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'ul', 'ol']:
            # Para blocos, pode ser melhor processar seus filhos diretamente
            # ou obter o texto e adicioná-lo, dependendo da estrutura desejada.
            # Esta implementação tenta adicionar o texto do bloco com espaços.
            text_from_block = child.get_text(separator=" ", strip=True)
            if text_from_block:
                paragraph.add_run(" " + text_from_block + " ")  # Adiciona espaços para separar de outros runs
        else:  # Trata outros elementos inline
            text_from_inline = child.get_text(strip=True)
            if text_from_inline:
                paragraph.add_run(text_from_inline)


def create_docx_from_text_or_html(content_input, is_html=False, title="Resposta do Chat"):
    document = Document()
    document.add_heading(title, level=1)
    bio = BytesIO()
    try:
        if is_html:
            soup = BeautifulSoup(content_input, 'html.parser')
            # Processa elementos de forma mais granular para melhor formatação
            for element in soup.find_all(True, recursive=False):  # Pega todos os elementos de nível superior
                if element.name.startswith('h') and len(element.name) == 2 and element.name[1].isdigit():
                    level = int(element.name[1])
                    heading_paragraph = document.add_heading(level=min(level, 9))  # Docx suporta até nível 9
                    add_runs_from_html_element(heading_paragraph, element)
                elif element.name in ['p', 'div']:
                    p = document.add_paragraph()
                    add_runs_from_html_element(p, element)
                elif element.name in ['ul', 'ol']:
                    list_style = 'ListBullet' if element.name == 'ul' else 'ListNumber'
                    for li in element.find_all('li', recursive=False):  # Apenas <li> diretos
                        item_p = document.add_paragraph(style=list_style)
                        add_runs_from_html_element(item_p, li)
                # Adicione mais manipulação para outros elementos HTML se necessário (ex: tabelas)
                else:  # Se for um elemento não tratado especificamente, mas tem texto
                    text_content = element.get_text(separator=" ", strip=True)
                    if text_content:
                        p = document.add_paragraph()
                        add_runs_from_html_element(p, element)  # Tenta processar seus filhos

            # Fallback se nenhum elemento de bloco principal foi encontrado, mas há texto
            if not soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'div'],
                                 recursive=False) and soup.get_text(strip=True):
                p = document.add_paragraph()
                add_runs_from_html_element(p, soup)  # Processa o soup como um todo

        else:  # Conteúdo é texto simples
            for line in content_input.split('\n'):
                document.add_paragraph(line)

        document.save(bio)
        bio.seek(0)
        return bio
    except Exception as e:
        # Fallback para documento de erro
        error_doc = Document()
        error_doc.add_heading("Erro na Conversão para DOCX", level=1)
        error_doc.add_paragraph(f"Ocorreu um erro ao tentar converter o conteúdo para DOCX.")
        error_doc.add_paragraph(f"Detalhes do erro: {str(e)}")
        error_doc.add_heading("Conteúdo Original (ou parte dele):", level=2)
        content_str = str(content_input) if content_input is not None else "[Conteúdo Nulo]"
        max_len = 5000  # Limita o tamanho do conteúdo no docx de erro
        content_to_add = content_str[:max_len] + "\n... (conteúdo truncado)" if len(
            content_str) > max_len else content_str
        error_doc.add_paragraph(content_to_add)
        error_bio_fallback = BytesIO()
        error_doc.save(error_bio_fallback)
        error_bio_fallback.seek(0)
        return error_bio_fallback


# <<< FUNÇÃO NOVA para extrair texto de arquivos >>>
def extract_text_from_file(uploaded_file):
    """Extrai texto de um arquivo carregado (txt, pdf, docx)."""
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    text_content = ""
    try:
        if file_extension == ".txt":
            text_content = uploaded_file.read().decode("utf-8", errors="ignore")
        elif file_extension == ".pdf":
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
        elif file_extension == ".docx":
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        else:
            return None, f"Formato não suportado: {uploaded_file.name}"
        return text_content.strip(), None
    except Exception as e:
        return None, f"Erro ao processar '{uploaded_file.name}': {str(e)}"


# --- Funções de API ---
def transcribe_with_groq(api_key, audio_file_bytes, original_filename):
    if not api_key:
        st.error("Chave API da Groq não configurada em .streamlit/secrets.toml. Necessária para transcrição.")
        return None  # Modificado para retornar None explicitamente

    headers = {"Authorization": f"Bearer {api_key}"}
    files = {"file": (original_filename, audio_file_bytes, "audio/mpeg")}
    data = {"model": SELECTED_TRANSCRIPTION_MODEL, "language": "pt"}

    try:
        response = requests.post(GROQ_API_TRANSCRIPTIONS_ENDPOINT, headers=headers, files=files, data=data)
        response.raise_for_status()
        try:
            return response.json()["text"]
        except requests.exceptions.JSONDecodeError:
            return response.text
    except requests.exceptions.HTTPError as http_err:
        st.error(f"Transcrição ({original_filename}) - Erro HTTP: {http_err} - {response.text}")
    except requests.exceptions.RequestException as req_err:
        st.error(f"Transcrição ({original_filename}) - Erro na requisição: {req_err}")
    except Exception as e:
        st.error(f"Transcrição ({original_filename}) - Erro inesperado: {e}")
    return None


def query_chatvolt_agent(api_key, agent_id, query, conversation_id=None, visitor_id=None):
    if not api_key or not agent_id:
        st.error("Chatvolt - Chave API ou ID do Agente não configurados em .streamlit/secrets.toml.")
        return None
    url = f"{CHATVOLT_API_BASE_URL}/{agent_id}/query"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    data = {"query": query, "streaming": False}
    if conversation_id:
        data["conversationId"] = conversation_id
    if visitor_id:
        data["visitorId"] = visitor_id
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.HTTPError as http_err:
        st.error(f"Chatvolt - Erro HTTP: {http_err} - {response.text}")
        return None
    except Exception as err:
        st.error(f"Chatvolt - Outro erro: {err}")
        return None


@st.cache_data(ttl=3600)
def get_groq_models(api_key):
    if not api_key:
        # Não mostra erro aqui, pois a UI da sidebar informará
        return []
    url = f"{GROQ_API_BASE_URL}/models"
    headers = {"Authorization": f"Bearer {api_key}"}
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        models_data = response.json()
        available_models = [model['id'] for model in models_data.get('data', []) if model.get('id')]
        # Prioriza modelos comuns se estiverem disponíveis
        priority_models = [m for m in COMMON_GROQ_MODELS if m in available_models]
        other_models = sorted([m for m in available_models if m not in COMMON_GROQ_MODELS])
        final_list = priority_models + other_models
        return final_list if final_list else sorted(
            list(set(available_models)))  # Garante que algo seja retornado se a lógica de prioridade falhar
    except requests.exceptions.HTTPError:
        # st.warning("Não foi possível buscar modelos Groq devido a um erro HTTP. Verifique sua chave API e conexão.")
        return []  # Retorna lista vazia em caso de erro HTTP (chave inválida, etc)
    except Exception:
        # st.warning("Ocorreu um erro inesperado ao buscar modelos Groq.")
        return []  # Retorna lista vazia para outros erros


def query_groq_api(api_key, model_id, messages_history):
    if not api_key or not model_id:
        st.error("Groq - Chave API não configurada em .streamlit/secrets.toml ou Modelo não selecionado.")
        return None
    url = f"{GROQ_API_BASE_URL}/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    data = {"model": model_id, "messages": messages_history, "temperature": 0.7}
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.HTTPError as http_err:
        st.error(f"Groq - Erro HTTP: {http_err} - {response.text}")
        return None
    except Exception as err:
        st.error(f"Groq - Outro erro: {err}")
        return None


# --- Gerenciamento de Estado e Navegação ---
def initialize_session_state():
    defaults = {
        "current_page": "input_fatos",
        "fatos_text": "",
        "fatos_text_buffer": st.session_state.get("fatos_text_buffer", ""),
        "selected_chat_type": None,
        "initial_prompt_processed": False,
        "chatvolt_messages": [],
        "chatvolt_conversation_id": None,
        "chatvolt_visitor_id": None,
        "groq_messages": [],
        # Chaves API e Agent ID não são mais armazenadas no session_state globalmente,
        # serão lidas de st.secrets e passadas via app_configs
        "selected_groq_model_global": st.session_state.get("selected_groq_model_global", None),
         # NOVOS ESTADOS PARA BUSCA DE JURISPRUDÊNCIA
        "termo_jurisprudencia": "",
        "resultados_jurisprudencia": None,
        "buscando_jurisprudencia": False  # Para controlar o spinner e a lógica de busca
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def navigate_to(page_name):
    st.session_state.current_page = page_name


def reset_all_chat_states():
    st.session_state.selected_chat_type = None
    st.session_state.initial_prompt_processed = False
    st.session_state.chatvolt_messages = []
    st.session_state.chatvolt_conversation_id = None
    st.session_state.chatvolt_visitor_id = None
    st.session_state.groq_messages = []


def reset_for_new_fatos():
    st.session_state.fatos_text = ""
    st.session_state.fatos_text_buffer = ""
    reset_all_chat_states()
    navigate_to("input_fatos")
    st.rerun()


# --- Componentes da Interface (UI) ---
# ... (outras partes do código) ...
def render_sidebar(available_groq_models):
    with st.sidebar:
        st.header("🔑 Configurações API")

        groq_api_key = st.secrets.get("groq_api_key")
        chatvolt_api_key = st.secrets.get("chatvolt_api_key")
        chatvolt_agent_id = st.secrets.get("chatvolt_agent_id")

        if groq_api_key:
            st.success("Chave API Groq: Carregada") # Removido `secrets.toml` para brevidade
        else:
            st.error("Chave API Groq: Não encontrada")
            st.caption("Adicione `groq_api_key` a `.streamlit/secrets.toml`")

        st.markdown("---")
        st.header("Chatvolt")
        if chatvolt_api_key:
            st.success("Chave API Chatvolt: Carregada")
        else:
            st.error("Chave API Chatvolt: Não encontrada")
            st.caption("Adicione `chatvolt_api_key` a `.streamlit/secrets.toml`")

        if chatvolt_agent_id:
            st.success(f"ID Agente Chatvolt: Carregado") # Removido trecho do ID
        else:
            st.error("ID Agente Chatvolt: Não encontrado")
            st.caption("Adicione `chatvolt_agent_id` a `.streamlit/secrets.toml`")

        st.markdown("---")
        st.header("Modelos Groq")
        # ... (lógica dos modelos Groq existente) ...
        selected_model_session = st.session_state.get("selected_groq_model_global") #
        current_model_index = 0 #

        if available_groq_models: #
            if selected_model_session and selected_model_session in available_groq_models: #
                current_model_index = available_groq_models.index(selected_model_session) #
            elif COMMON_GROQ_MODELS[0] in available_groq_models: #
                current_model_index = available_groq_models.index(COMMON_GROQ_MODELS[0]) #
            st.session_state.selected_groq_model_global = st.selectbox( #
                "Escolha o modelo Groq para chat:", #
                options=available_groq_models, #
                index=current_model_index, #
                key="sb_groq_model_selector" #
            )
        elif groq_api_key: #
            st.warning("Não foi possível carregar modelos Groq. Verifique a chave API ou a conexão.") #
            st.session_state.selected_groq_model_global = None #
        else: #
            st.info("Modelos Groq aparecerão aqui após configurar a chave API.") #
            st.session_state.selected_groq_model_global = None #


        st.markdown("---")
        st.header("Navegação Principal") # Novo subcabeçalho para clareza
        if st.button("Registrar Novos Fatos", key="reset_sidebar_button", help="Limpar dados atuais e registrar novos fatos."): # Adicionado help
            reset_for_new_fatos()
        
        if st.button("⚖️ Buscar Jurisprudência (TJGO)", key="btn_to_jurisprudencia_search", help="Pesquisar na base de jurisprudência do TJGO."): # Adicionado help
            navigate_to("busca_jurisprudencia")
            st.rerun() # Este rerun para navegação é geralmente OK

        st.markdown("---")
        st.caption("Assistente Jurídico v1.0") # Apenas um caption no final

    return {
        "groq_api_key": groq_api_key,
        "chatvolt_api_key": chatvolt_api_key,
        "chatvolt_agent_id": chatvolt_agent_id,
        "selected_groq_model": st.session_state.selected_groq_model_global
    }
# ... (restante do código) ...
    with st.sidebar:
        st.header("🔑 Configurações API")

        # Carregar chaves do secrets.toml
        groq_api_key = st.secrets.get("groq_api_key")
        chatvolt_api_key = st.secrets.get("chatvolt_api_key")
        chatvolt_agent_id = st.secrets.get("chatvolt_agent_id")

        if groq_api_key:
            st.success("Chave API Groq: Carregada de `secrets.toml`")
        else:
            st.error("Chave API Groq: Não encontrada em `secrets.toml`")
            st.caption("Adicione `groq_api_key = \"SUA_CHAVE\"` ao arquivo `.streamlit/secrets.toml`")

        st.markdown("---")
        st.header("Chatvolt")
        if chatvolt_api_key:
            st.success("Chave API Chatvolt: Carregada de `secrets.toml`")
        else:
            st.error("Chave API Chatvolt: Não encontrada em `secrets.toml`")
            st.caption("Adicione `chatvolt_api_key = \"SUA_CHAVE\"` ao arquivo `.streamlit/secrets.toml`")

        if chatvolt_agent_id:
            st.success(f"ID Agente Chatvolt: Carregado (`{chatvolt_agent_id[:4]}...`)")
        else:
            st.error("ID Agente Chatvolt: Não encontrado em `secrets.toml`")
            st.caption("Adicione `chatvolt_agent_id = \"SEU_ID\"` ao arquivo `.streamlit/secrets.toml`")

        st.markdown("---")
        st.header("Modelos Groq")

        selected_model_session = st.session_state.get("selected_groq_model_global")
        current_model_index = 0

        if available_groq_models:
            if selected_model_session and selected_model_session in available_groq_models:
                current_model_index = available_groq_models.index(selected_model_session)
            elif COMMON_GROQ_MODELS[0] in available_groq_models:  # Tenta o primeiro modelo comum como padrão
                current_model_index = available_groq_models.index(COMMON_GROQ_MODELS[0])
            # Se nenhum modelo estiver selecionado ou o anterior não for válido, usa o índice 0 (o primeiro da lista)

            st.session_state.selected_groq_model_global = st.selectbox(
                "Escolha o modelo Groq para chat:",
                options=available_groq_models,
                index=current_model_index,
                key="sb_groq_model_selector"
            )
        elif groq_api_key:  # Se a chave existe mas não carregou modelos
            st.warning("Não foi possível carregar modelos Groq. Verifique a chave API ou a conexão.")
            st.session_state.selected_groq_model_global = None
        else:  # Se a chave não existe
            st.info("Modelos Groq aparecerão aqui após configurar a chave API.")
            st.session_state.selected_groq_model_global = None

        st.markdown("---")
        if st.button("Registrar Novos Fatos", key="reset_sidebar_button"):
            reset_for_new_fatos()
        st.markdown("---")
        st.caption("Assistente Jurídico")
                # NOVO BOTÃO PARA BUSCA DE JURISPRUDÊNCIA
        if st.button("⚖️ Buscar Jurisprudência (TJGO)", key="btn_to_jurisprudencia_search"):
            navigate_to("busca_jurisprudencia")
            st.rerun()
        st.markdown("---")
        st.caption("Assistente Jurídico")
    # ... (resto do código da sidebar) ...

    return {
        "groq_api_key": groq_api_key,
        "chatvolt_api_key": chatvolt_api_key,
        "chatvolt_agent_id": chatvolt_agent_id,
        "selected_groq_model": st.session_state.selected_groq_model_global
    }

# interface.py
# ... (seu código existente) ...

# No arquivo interface.py
# ... (outras partes do código) ...

def render_busca_jurisprudencia_page(app_configs):
    st.title("⚖️ Busca de Jurisprudência - TJGO")
    st.markdown("Insira o termo que deseja pesquisar na base de jurisprudência do TJGO.")

    termo_busca_input = st.text_input(
        "Termo de busca:",
        value=st.session_state.get("termo_jurisprudencia", ""),
        key="termo_jurisprudencia_input_key"
    )
    if termo_busca_input != st.session_state.get("termo_jurisprudencia"):
        st.session_state.termo_jurisprudencia = termo_busca_input
        # Um rerun aqui pode ser desnecessário se o botão de busca for a ação principal
        # st.rerun()

    if st.button("Buscar Jurisprudência", key="btn_buscar_jurisprudencia_action"):
        if not st.session_state.termo_jurisprudencia.strip():
            st.warning("Por favor, insira um termo para a busca.")
        else:
            # Não fazer rerun aqui. Deixar o spinner controlar a próxima renderização.
            st.session_state.buscando_jurisprudencia = True
            st.session_state.resultados_jurisprudencia = None
            # st.rerun() # REMOVER ESTE RERUN

    if st.session_state.get("buscando_jurisprudencia"):
        termo_para_busca = st.session_state.termo_jurisprudencia
        # Usar st.status para uma melhor experiência com o spinner
        with st.status(f"Buscando jurisprudência para: '{termo_para_busca}'...", expanded=True) as status_ui:
            try:
                st.write(f"Iniciando busca no TJGO para: {termo_para_busca}")
                script_path = os.path.join(os.path.dirname(__file__), 'jurisprudencia.py')
                if not os.path.exists(script_path):
                    st.error(f"Script 'jurisprudencia.py' não encontrado em: {script_path}")
                    st.session_state.resultados_jurisprudencia = [{"erro_interno": "jurisprudencia.py não encontrado."}]
                else:
                    process = subprocess.run(
                        [sys.executable, script_path, termo_para_busca],
                        capture_output=True, text=True, check=False, encoding='utf-8', timeout=120
                    )
                    if process.returncode != 0:
                        st.error(f"Script de busca falhou. Erro: {process.stderr}")
                        st.session_state.resultados_jurisprudencia = [{"erro_subprocess": f"Erro script: {process.stderr}"}]
                    else:
                        resultados_raw = process.stdout
                        try:
                            resultados_json = json.loads(resultados_raw)
                            st.session_state.resultados_jurisprudencia = resultados_json
                            st.write("Busca concluída. Processando resultados...")
                        except json.JSONDecodeError:
                            st.error(f"Erro ao decodificar JSON do script: {resultados_raw}")
                            st.session_state.resultados_jurisprudencia = [{"erro_json_decode": f"Falha JSON: {resultados_raw}"}]
                status_ui.update(label="Busca finalizada!", state="complete")

            except subprocess.TimeoutExpired:
                st.error("A busca de jurisprudência demorou muito (timeout).")
                st.session_state.resultados_jurisprudencia = [{"erro_timeout": "Busca excedeu o tempo limite."}]
                status_ui.update(label="Timeout na busca!", state="error")
            except FileNotFoundError: # Redundante se o check os.path.exists for feito
                st.error("Erro: Script 'jurisprudencia.py' não encontrado.")
                st.session_state.resultados_jurisprudencia = [{"erro_interno": "jurisprudencia.py não encontrado."}]
                status_ui.update(label="Erro de arquivo!", state="error")
            except Exception as e:
                st.error(f"Erro inesperado na busca: {str(e)}")
                st.session_state.resultados_jurisprudencia = [{"erro_inesperado": str(e)}]
                status_ui.update(label="Erro inesperado!", state="error")
            finally:
                st.session_state.buscando_jurisprudencia = False
                # O st.rerun() AQUI É O MAIS CRÍTICO A SER CONSIDERADO.
                # Se a atualização do session_state for suficiente para o Streamlit
                # re-renderizar a exibição dos resultados, este rerun pode não ser necessário
                # ou pode ser a causa do problema "removeChild".
                # Tente comentar este rerun primeiro. Se a UI não atualizar
                # com os resultados, então ele pode ser necessário, mas pode precisar
                # de uma lógica mais cuidadosa.
                # st.rerun() # <= TENTE COMENTAR ESTE PRIMEIRO

    # Exibe os resultados após a busca (esta parte permanece a mesma)
    if not st.session_state.get("buscando_jurisprudencia") and st.session_state.get("resultados_jurisprudencia") is not None:
        resultados = st.session_state.get("resultados_jurisprudencia")
        st.subheader("Resultados da Busca:")
        if isinstance(resultados, list) and resultados:
            # ... (lógica de exibição de resultados existente) ...
            has_actual_results = False
            for i, res in enumerate(resultados):
                error_keys = ["erro_driver", "erro_geral", "erro_subprocess", "erro_json_decode", "erro_interno", "erro_inesperado", "erro_timeout"]
                found_error = False
                for key_err in error_keys: # Renomear variável de loop para evitar conflito
                    if key_err in res:
                        st.error(f"Erro na busca: {res[key_err]}")
                        if key_err == "erro_driver": st.info("Verifique Google Chrome / ChromeDriver.")
                        found_error = True
                        break
                if found_error: continue

                if "info" in res:
                    st.info(res["info"])
                    continue

                has_actual_results = True
                st.markdown(f"---")
                with st.container(border=True): # 'border' é um bom parâmetro
                    st.markdown(f"**Resultado {res.get('id', i+1)}**")
                    if "texto" in res and res["texto"]:
                        st.text_area(f"Conteúdo do Resultado {res.get('id', i+1)}:", value=res["texto"], height=200, key=f"juris_text_display_{res.get('id', i)}", disabled=True)
                    elif "erro" in res:
                        st.warning(f"Falha ao obter conteúdo do resultado {res.get('id', i+1)}: {res['erro']}")
                    else:
                        st.warning(f"Resultado {res.get('id', i+1)} em formato inesperado ou sem conteúdo.")
            if not has_actual_results and not any("info" in r for r in resultados if isinstance(r, dict)):
                 st.info("A busca foi concluída, mas não retornou jurisprudências ou houve apenas mensagens de erro.")

        elif not resultados:
             st.info("A busca não retornou dados.")
        else:
            st.warning("Formato de resultados da busca inesperado.")

    st.markdown("---")
    if st.button("Voltar para Registro de Fatos", key="btn_juris_to_fatos"):
        st.session_state.termo_jurisprudencia = ""
        st.session_state.resultados_jurisprudencia = None
        st.session_state.buscando_jurisprudencia = False
        navigate_to("input_fatos")
        st.rerun() # Este rerun para navegação é geralmente OK.
# ... (restante do código) ...    st.title("⚖️ Busca de Jurisprudência - TJGO")
    st.markdown("Insira o termo que deseja pesquisar na base de jurisprudência do TJGO.")

# interface.py, aproximadamente na linha 572
def render_busca_jurisprudencia_page(app_configs):
    st.title("⚖️ Busca de Jurisprudência - TJGO")
    st.markdown("Insira o termo que deseja pesquisar na base de jurisprudência do TJGO.")

    termo_busca_input = st.text_input( # ESTE É O WIDGET CAUSANDO O ERRO
        "Termo de busca:",
        value=st.session_state.get("termo_jurisprudencia", ""),
        key="termo_jurisprudencia_input_key" # ESTA CHAVE ESTÁ DUPLICADA
    )
    # ... restante da função ...
    # Atualizar o estado da sessão se o valor do input mudar
    if termo_busca_input != st.session_state.get("termo_jurisprudencia"):
        st.session_state.termo_jurisprudencia = termo_busca_input
        # st.rerun() # Pode não ser necessário aqui, depende da interatividade desejada

    if st.button("Buscar Jurisprudência", key="btn_buscar_jurisprudencia_action"):
        if not st.session_state.termo_jurisprudencia.strip():
            st.warning("Por favor, insira um termo para a busca.")
        else:
            st.session_state.buscando_jurisprudencia = True
            st.session_state.resultados_jurisprudencia = None # Limpa resultados anteriores
            st.rerun() # Para mostrar o spinner imediatamente

    if st.session_state.get("buscando_jurisprudencia"):
        termo_para_busca = st.session_state.termo_jurisprudencia
        with st.spinner(f"Buscando jurisprudência para: '{termo_para_busca}'... Aguarde, isso pode levar alguns instantes."):
            try:
                # Certifique-se que 'jurisprudencia.py' está no mesmo diretório
                # Usar sys.executable para garantir que o mesmo interpretador Python seja usado
                process = subprocess.run(
                    [sys.executable, 'jurisprudencia.py', termo_para_busca],
                    capture_output=True,
                    text=True,
                    check=True, # Lança CalledProcessError se o script retornar um código de erro
                    encoding='utf-8' # Importante para caracteres especiais
                )
                resultados_raw = process.stdout
                try:
                    resultados_json = json.loads(resultados_raw)
                    st.session_state.resultados_jurisprudencia = resultados_json
                except json.JSONDecodeError:
                    st.error(f"Erro ao decodificar a resposta JSON do script de jurisprudência. Resposta recebida:\n{resultados_raw}")
                    st.session_state.resultados_jurisprudencia = [{"erro_json_decode": f"Falha na decodificação JSON. Resposta: {resultados_raw}"}]

            except subprocess.CalledProcessError as e:
                st.error("Ocorreu um erro ao executar a busca de jurisprudência.")
                st.error(f"Detalhes do erro do script:\n{e.stderr}")
                st.session_state.resultados_jurisprudencia = [{"erro_subprocess": f"Erro no script: {e.stderr}"}]
            except FileNotFoundError:
                st.error("Erro: O script 'jurisprudencia.py' não foi encontrado. Certifique-se de que ele está no mesmo diretório que esta aplicação.")
                st.session_state.resultados_jurisprudencia = [{"erro_interno": "Script jurisprudencia.py não encontrado."}]
            except Exception as e:
                st.error(f"Um erro inesperado ocorreu durante a busca: {str(e)}")
                st.session_state.resultados_jurisprudencia = [{"erro_inesperado": str(e)}]
            finally:
                st.session_state.buscando_jurisprudencia = False
                st.rerun() # Para exibir os resultados ou erros e remover o spinner

    # Exibe os resultados após a busca
    if not st.session_state.get("buscando_jurisprudencia") and st.session_state.get("resultados_jurisprudencia") is not None:
        resultados = st.session_state.get("resultados_jurisprudencia")
        st.subheader("Resultados da Busca:")
        if isinstance(resultados, list) and resultados:
            for i, res in enumerate(resultados):
                # Verifica os tipos de erro primeiro
                if "erro_driver" in res:
                    st.error(f"Erro Crítico na Busca (WebDriver): {res['erro_driver']}")
                    st.info("Verifique se o Google Chrome está instalado e se não há problemas com o ChromeDriver.")
                    break
                if "erro_geral" in res:
                    st.error(f"Erro Geral na Busca: {res['erro_geral']}")
                    break
                if "erro_subprocess" in res:
                    st.error(f"Erro na Execução do Script de Busca: {res['erro_subprocess']}")
                    break
                if "erro_json_decode" in res:
                    st.error(f"Erro Interno (JSON Decode): {res['erro_json_decode']}")
                    break
                if "erro_interno" in res: # Erros como FileNotFoundError
                    st.error(f"Erro Interno: {res['erro_interno']}")
                    break
                if "erro_inesperado" in res:
                    st.error(f"Erro Inesperado: {res['erro_inesperado']}")
                    break
                if "info" in res: # Mensagens informativas como "nenhum resultado"
                    st.info(res["info"])
                    break

                # Se chegou aqui, é um resultado válido ou um erro de processamento de bloco
                st.markdown(f"--- **Resultado {res.get('id', i+1)}** ---")
                if "texto" in res and res["texto"]:
                    st.text_area(f"Jurisprudência {res.get('id', i+1)}:", value=res["texto"], height=250, key=f"juris_text_{i}", disabled=True)
                elif "erro" in res: # Erro específico ao processar um bloco
                    st.warning(f"Falha ao processar o conteúdo do resultado {res.get('id', i+1)}: {res['erro']}")
                else: # Caso algum resultado venha em formato inesperado, sem 'texto' ou 'erro'
                    st.warning(f"Resultado {res.get('id', i+1)} em formato inesperado ou sem conteúdo.")
        elif not resultados: # Lista vazia, pode acontecer se o script retornar []
             st.info("A busca não retornou resultados ou a lista de resultados está vazia.")
        else: # Não é lista ou é None (embora a lógica acima deva cobrir None)
            st.warning("Formato de resultados da busca inesperado.")

    st.markdown("---")
    if st.button("Voltar para Registro de Fatos", key="btn_juris_to_fatos"):
        # Limpar estados da página de jurisprudência ao sair
        st.session_state.termo_jurisprudencia = ""
        st.session_state.resultados_jurisprudencia = None
        st.session_state.buscando_jurisprudencia = False
        navigate_to("input_fatos")
        st.rerun()
# ... (seu código existente) ...
def render_fatos_input_page(app_configs):
    st.title("📝 Registro de Fatos")
    st.markdown("Descreva os fatos ou utilize as opções abaixo para transcrever áudios e anexar documentos.")
    st.markdown("---")

    st.subheader("🎙️ Transcrever Áudio para Fatos (Opcional)")
    uploaded_audio_files = st.file_uploader(
        "Envie um ou mais arquivos de áudio (MP3, WAV, M4A, OGG, OPUS, etc.):",
        type=["mp3", "wav", "m4a", "ogg", "opus", "mp4", "mpeg", "mpga", "webm"],
        accept_multiple_files=True,
        key="fatos_audio_uploader"
    )

    if uploaded_audio_files:
        if st.button("➕ Adicionar Transcrição(ões) aos Fatos", key="btn_transcribe_fatos"):
            groq_api_key = app_configs.get("groq_api_key")  # Vem de st.secrets via app_configs
            if not groq_api_key:
                st.error("Chave API da Groq não configurada em `.streamlit/secrets.toml`. Necessária para transcrição.")
            else:
                all_transcriptions_texts = []
                has_errors_in_transcription = False
                progress_bar = st.progress(0)
                status_text = st.empty()

                for i, audio_file in enumerate(uploaded_audio_files):
                    current_progress = (i + 1) / len(uploaded_audio_files)
                    status_text.info(
                        f"Processando áudio {i + 1} de {len(uploaded_audio_files)}: '{audio_file.name}'...")

                    file_size_mb = audio_file.size / (1024 * 1024)
                    if file_size_mb > MAX_AUDIO_FILE_SIZE_MB:
                        st.warning(
                            f"Áudio '{audio_file.name}' ({file_size_mb:.2f}MB) excede o limite de {MAX_AUDIO_FILE_SIZE_MB}MB e será ignorado.")
                        all_transcriptions_texts.append(
                            f"\n--- [Áudio '{audio_file.name}' ignorado: tamanho excede o limite] ---\n")
                        has_errors_in_transcription = True
                        progress_bar.progress(current_progress)
                        continue

                    audio_bytes = audio_file.getvalue()
                    # Passa a chave API diretamente para a função de transcrição
                    transcription = transcribe_with_groq(groq_api_key, audio_bytes, audio_file.name)

                    if transcription:
                        all_transcriptions_texts.append(
                            f"\n--- Transcrição de '{audio_file.name}' ---\n{transcription}\n--- Fim da Transcrição de '{audio_file.name}' ---")
                    else:
                        all_transcriptions_texts.append(f"\n--- [Falha na transcrição de '{audio_file.name}'] ---")
                        has_errors_in_transcription = True
                    progress_bar.progress(current_progress)

                status_text.empty()
                progress_bar.empty()

                if all_transcriptions_texts:
                    current_buffer = st.session_state.fatos_text_buffer
                    new_text = "\n\n".join(all_transcriptions_texts)
                    if current_buffer.strip():
                        st.session_state.fatos_text_buffer = f"{current_buffer}\n\n{new_text}"
                    else:
                        st.session_state.fatos_text_buffer = new_text.strip()
                    st.rerun()

                if not has_errors_in_transcription and all_transcriptions_texts:
                    st.success("Transcrição(ões) adicionada(s) ao campo 'Descrição dos Fatos'. Revise e edite.")
                elif all_transcriptions_texts:
                    st.info("Processamento de áudio concluído. Verifique os resultados no campo 'Descrição dos Fatos'.")
    st.markdown("---")

    st.subheader("📄 Anexar Arquivos de Texto (Opcional)")
    uploaded_text_files = st.file_uploader(
        "Envie um ou mais arquivos de texto (TXT, PDF, DOCX):",
        type=ALLOWED_TEXT_EXTENSIONS,
        accept_multiple_files=True,
        key="fatos_text_file_uploader"
    )

    if uploaded_text_files:
        if st.button("➕ Adicionar Conteúdo do(s) Arquivo(s) aos Fatos", key="btn_add_text_files"):
            all_extracted_texts = []
            has_errors_in_extraction = False
            text_file_progress_bar = st.progress(0)
            text_file_status_text = st.empty()

            for i, text_file in enumerate(uploaded_text_files):
                current_progress = (i + 1) / len(uploaded_text_files)
                text_file_status_text.info(
                    f"Processando arquivo de texto {i + 1} de {len(uploaded_text_files)}: '{text_file.name}'...")
                extracted_content, error_msg = extract_text_from_file(text_file)
                if error_msg:
                    st.warning(f"Arquivo '{text_file.name}': {error_msg}")
                    all_extracted_texts.append(f"\n--- [Falha ao ler o arquivo '{text_file.name}': {error_msg}] ---")
                    has_errors_in_extraction = True
                elif extracted_content:
                    all_extracted_texts.append(
                        f"\n--- Conteúdo de '{text_file.name}' ---\n{extracted_content}\n--- Fim do Conteúdo de '{text_file.name}' ---")
                else:
                    all_extracted_texts.append(
                        f"\n--- [Arquivo '{text_file.name}' não continha texto extraível ou estava vazio] ---")
                text_file_progress_bar.progress(current_progress)

            text_file_status_text.empty()
            text_file_progress_bar.empty()

            if all_extracted_texts:
                current_buffer = st.session_state.fatos_text_buffer
                new_text_from_files = "\n\n".join(all_extracted_texts)
                if current_buffer.strip():
                    st.session_state.fatos_text_buffer = f"{current_buffer}\n\n{new_text_from_files}"
                else:
                    st.session_state.fatos_text_buffer = new_text_from_files.strip()
                st.rerun()

            if not has_errors_in_extraction and all_extracted_texts:
                st.success("Conteúdo do(s) arquivo(s) adicionado(s) ao campo 'Descrição dos Fatos'. Revise e edite.")
            elif all_extracted_texts:
                st.info(
                    "Processamento de arquivos de texto concluído. Verifique os resultados no campo 'Descrição dos Fatos'.")
    st.markdown("---")

    st.subheader("📝 Descrição dos Fatos")
    fatos_input_value = st.session_state.fatos_text_buffer
    edited_fatos_text = st.text_area(
        "Detalhe os fatos aqui:",
        value=fatos_input_value,
        height=400,
        max_chars=100000,
        key="fatos_input_area_ta_main"
    )
    if edited_fatos_text != fatos_input_value:
        st.session_state.fatos_text_buffer = edited_fatos_text
        # Não precisa de st.rerun() aqui, o widget já atualiza o valor no próximo ciclo
        # st.rerun() # Removido para evitar reruns desnecessários a cada digitação

    if st.button("Prosseguir para Seleção do Assistente", key="btn_to_select_chat"):
        if st.session_state.fatos_text_buffer.strip():
            st.session_state.fatos_text = st.session_state.fatos_text_buffer.strip()
            # Resetar estados de chat antes de ir para a seleção
            reset_all_chat_states()  # Garante que estados de chat anteriores sejam limpos
            navigate_to("select_chat")
            st.rerun()
        else:
            st.warning("Por favor, descreva os fatos antes de prosseguir.")


def render_chat_selection_page(app_configs):
    st.title("🤖 Escolha o Assistente")
    st.markdown(f"""
    **Fatos Registrados (Prévia):**
    ```
    {st.session_state.fatos_text[:300]}{'...' if len(st.session_state.fatos_text) > 300 else ''}
    ```
    Selecione qual assistente você gostaria de usar para analisar estes fatos.
    """)

    col1_select, col2_select = st.columns(2)
    with col1_select:
        if st.button("⚖️ Usar Assistente Jurídico Principal (Chatvolt)", key="btn_use_chatvolt",
                     use_container_width=True):
            if not app_configs["chatvolt_api_key"] or not app_configs["chatvolt_agent_id"]:
                st.error("Chatvolt não configurado. Verifique `secrets.toml` na pasta `.streamlit`.")
            else:
                st.session_state.selected_chat_type = "chatvolt"
                navigate_to("chat_view")
                st.rerun()
    with col2_select:
        if st.button("🧠 Usar Assistente Geral Rápido (Groq)", key="btn_use_groq",
                     use_container_width=True):
            if not app_configs["groq_api_key"]:
                st.error("Groq não configurado. Verifique `secrets.toml` na pasta `.streamlit`.")
            elif not app_configs["selected_groq_model"]:
                st.error("Nenhum modelo Groq selecionado ou disponível. Verifique as configurações na barra lateral.")
            else:
                st.session_state.selected_chat_type = "groq"
                navigate_to("chat_view")
                st.rerun()

    if st.button("Voltar e Editar Fatos", key="btn_back_to_fatos"):
        navigate_to("input_fatos")
        st.rerun()


def _handle_initial_prompt_processing(app_configs):
    chat_type = st.session_state.selected_chat_type
    chat_title_map = {"chatvolt": "Assistente Jurídico Principal", "groq": "Assistente Geral Rápido"}
    chat_title = chat_title_map.get(chat_type, "Assistente")

    with st.spinner(f"Analisando os fatos com {chat_title}..."):
        if chat_type == "chatvolt":
            # Adiciona a mensagem do usuário (fatos) antes de fazer a query
            st.session_state.chatvolt_messages.append({"role": "user", "content": st.session_state.fatos_text})
            response_data = query_chatvolt_agent(
                app_configs["chatvolt_api_key"], app_configs["chatvolt_agent_id"], st.session_state.fatos_text,
                st.session_state.chatvolt_conversation_id, st.session_state.chatvolt_visitor_id
            )
            assistant_response_text = "Desculpe, não consegui processar os fatos iniciais (Chatvolt)."
            docx_bytes = None
            msg_id = "cv_initial_error"
            sources = []

            if response_data:
                assistant_response_text = response_data.get("answer", "Não obtive uma resposta clara.")
                st.session_state.chatvolt_conversation_id = response_data.get("conversationId")
                st.session_state.chatvolt_visitor_id = response_data.get("visitorId")
                sources = response_data.get("sources", [])
                msg_id = response_data.get("messageId", "cv_initial_ok")

            docx_bytes = create_docx_from_text_or_html(assistant_response_text, is_html=True,
                                                       title=f"Resposta Inicial - {chat_title}")
            st.session_state.chatvolt_messages.append({
                "role": "assistant", "content": assistant_response_text,
                "sources": sources, "id": msg_id, "docx_bytes": docx_bytes
            })

        elif chat_type == "groq":
            # Adiciona a mensagem do usuário (fatos) antes de fazer a query
            st.session_state.groq_messages.append({"role": "user", "content": st.session_state.fatos_text})
            # Para Groq, o histórico completo de mensagens é normalmente enviado
            groq_history_for_api = [{"role": msg["role"], "content": msg["content"]} for msg in
                                    st.session_state.groq_messages]
            response_data = query_groq_api(
                app_configs["groq_api_key"], app_configs["selected_groq_model"], groq_history_for_api
            )
            assistant_response_text = "Desculpe, não consegui processar os fatos iniciais (Groq)."
            docx_bytes = None
            msg_id = "groq_initial_error"

            if response_data and response_data.get("choices"):
                assistant_response_text = response_data["choices"][0]["message"]["content"]
                msg_id = response_data.get("id", "groq_initial_ok")
            elif response_data and "error" in response_data:  # Trata erros da API Groq
                assistant_response_text += f" Detalhe: {response_data['error'].get('message', '')}"

            docx_bytes = create_docx_from_text_or_html(assistant_response_text, is_html=False,
                                                       # Groq geralmente não retorna HTML
                                                       title=f"Resposta Inicial - {chat_title}")
            st.session_state.groq_messages.append({
                "role": "assistant", "content": assistant_response_text,
                "id": msg_id, "docx_bytes": docx_bytes
            })

        st.session_state.initial_prompt_processed = True
        st.rerun()


def _display_chat_messages():
    chat_type = st.session_state.selected_chat_type
    current_messages = []
    if chat_type == "chatvolt":
        current_messages = st.session_state.chatvolt_messages
    elif chat_type == "groq":
        current_messages = st.session_state.groq_messages

    for i, message in enumerate(current_messages):
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            if message["role"] == "assistant":
                # Exibir fontes para Chatvolt se existirem
                if chat_type == "chatvolt" and "sources" in message and message["sources"]:
                    with st.expander("Ver fontes da resposta", expanded=False):  # Default para não expandido
                        for s_idx, source in enumerate(message["sources"]):
                            st.write(f"**Fonte {s_idx + 1}:** {source.get('text', 'N/A')}")
                            st.caption(
                                f"Documento: {source.get('datasource_name', 'N/A')}, Score: {source.get('score', 'N/A'):.2f}")
                            if source.get('document_url'):
                                st.link_button(f"Acessar Documento {s_idx + 1}", source['document_url'])
                            st.divider()

                if message.get("docx_bytes"):
                    file_name = f"resposta_{chat_type}_{message.get('id', f'msg{i}')}.docx"
                    st.download_button(
                        label="📥 Baixar Resposta (.docx)",
                        data=message["docx_bytes"],
                        file_name=file_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_btn_{chat_type}_{message.get('id', i)}"
                    )


def _handle_subsequent_user_input(app_configs, chat_title):
    if prompt := st.chat_input(f"Faça uma pergunta sobre os fatos para {chat_title}..."):
        chat_type = st.session_state.selected_chat_type
        user_message_data = {"role": "user", "content": prompt}

        current_message_list = None
        if chat_type == "chatvolt":
            current_message_list = st.session_state.chatvolt_messages
        elif chat_type == "groq":
            current_message_list = st.session_state.groq_messages

        if current_message_list is not None:
            current_message_list.append(user_message_data)

        with st.chat_message("user"):
            st.markdown(prompt)

        with st.spinner(f"{chat_title.split('(')[0].strip()} pensando..."):
            assistant_response_text = "Desculpe, não consegui processar sua solicitação."
            docx_bytes = None
            msg_id_suffix = len(current_message_list) if current_message_list is not None else "unk"
            msg_id = f"{chat_type}_subsequent_error_{msg_id_suffix}"

            if chat_type == "chatvolt":
                response_data = query_chatvolt_agent(
                    app_configs["chatvolt_api_key"], app_configs["chatvolt_agent_id"], prompt,
                    st.session_state.chatvolt_conversation_id, st.session_state.chatvolt_visitor_id
                )
                sources = []
                if response_data:
                    assistant_response_text = response_data.get("answer", "Não obtive uma resposta clara.")
                    st.session_state.chatvolt_conversation_id = response_data.get(
                        "conversationId")  # Atualiza ID da conversa
                    st.session_state.chatvolt_visitor_id = response_data.get("visitorId")  # Atualiza ID do visitante
                    sources = response_data.get("sources", [])
                    msg_id = response_data.get("messageId", f"cv_msg_{msg_id_suffix}")

                docx_bytes = create_docx_from_text_or_html(assistant_response_text, is_html=True,
                                                           title=f"Resposta - {chat_title}")
                st.session_state.chatvolt_messages.append({
                    "role": "assistant", "content": assistant_response_text,
                    "sources": sources, "id": msg_id, "docx_bytes": docx_bytes
                })

            elif chat_type == "groq":
                # Para Groq, o histórico completo de mensagens é normalmente enviado
                groq_history_for_api = [{"role": msg["role"], "content": msg["content"]} for msg in
                                        st.session_state.groq_messages]
                response_data = query_groq_api(app_configs["groq_api_key"], app_configs["selected_groq_model"],
                                               groq_history_for_api)

                if response_data and response_data.get("choices"):
                    assistant_response_text = response_data["choices"][0]["message"]["content"]
                    msg_id = response_data.get("id", f"groq_msg_{msg_id_suffix}")
                elif response_data and "error" in response_data:  # Trata erros da API Groq
                    assistant_response_text += f" Detalhe: {response_data['error'].get('message', '')}"

                docx_bytes = create_docx_from_text_or_html(assistant_response_text, is_html=False,
                                                           title=f"Resposta - {chat_title}")
                st.session_state.groq_messages.append({
                    "role": "assistant", "content": assistant_response_text,
                    "id": msg_id, "docx_bytes": docx_bytes
                })
        st.rerun()


def render_chat_view_page(app_configs):
    if not st.session_state.selected_chat_type:
        st.warning("Nenhum chat selecionado. Por favor, volte e escolha um assistente.")
        if st.button("Voltar para Seleção"):
            navigate_to("select_chat")
            st.rerun()
        st.stop()

    chat_title_map = {"chatvolt": "Assistente Jurídico Principal (Chatvolt)", "groq": "Assistente Geral Rápido (Groq)"}
    chat_title = chat_title_map.get(st.session_state.selected_chat_type, "Assistente")
    st.title(f"💬 {chat_title}")

    # Se os fatos ainda não foram processados, faz isso primeiro.
    if not st.session_state.initial_prompt_processed:
        _handle_initial_prompt_processing(app_configs)
        # O _handle_initial_prompt_processing já faz st.rerun() no final

    # Exibe todas as mensagens do chat
    _display_chat_messages()

    # Se o processamento inicial já ocorreu, permite novas entradas do usuário
    if st.session_state.initial_prompt_processed:
        _handle_subsequent_user_input(app_configs, chat_title)
    else:
        # Isso normalmente não deveria ser exibido se _handle_initial_prompt_processing funcionar
        st.info("Aguardando o processamento inicial dos fatos...")

    if st.button("Analisar Outros Fatos", key="btn_chat_to_fatos"):
        reset_for_new_fatos()  # Isso irá limpar estados e navegar para input_fatos


# --- Lógica Principal da Aplicação ---
def main():
    st.set_page_config(layout="wide", page_title="Assistente Jurídico")
    initialize_session_state()  # Garante que todos os estados de sessão necessários existam

    # Carrega a chave Groq de st.secrets para buscar modelos
    # Não precisamos armazenar isso em app_configs ainda, apenas para get_groq_models
    groq_api_key_for_models = st.secrets.get("groq_api_key")
    available_groq_models = get_groq_models(groq_api_key_for_models)

    # render_sidebar agora obtém as chaves de st.secrets internamente
    # e retorna um dicionário app_configs com essas chaves (ou None se não encontradas)
    # e o modelo Groq selecionado pelo usuário.
    app_configs = render_sidebar(available_groq_models)

    # Navegação entre páginas
    page_key = st.session_state.current_page
    if page_key == "input_fatos":
        render_fatos_input_page(app_configs)
    elif page_key == "select_chat":
        render_chat_selection_page(app_configs)  # Passa app_configs para verificar se as chaves estão carregadas
    elif page_key == "chat_view":
        render_chat_view_page(app_configs)  # Passa app_configs para uso nas chamadas de API
    # NOVA ROTA PARA BUSCA DE JURISPRUDÊNCIA
    elif page_key == "busca_jurisprudencia":
        render_busca_jurisprudencia_page(app_configs) # Passando app_configs por consistência
    else:
        st.error("Página desconhecida.")
        navigate_to("input_fatos")  # Volta para a página inicial em caso de erro
        st.rerun()


if __name__ == "__main__":
    main()
